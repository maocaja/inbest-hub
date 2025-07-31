import PyPDF2
import docx
import pandas as pd
import openpyxl
from fastapi import UploadFile
import io
import logging
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Servicio para procesar documentos PDF, DOCX y Excel
    """
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel
        }
    
    async def process_document(self, file: UploadFile) -> str:
        """
        Procesa un documento y extrae su contenido como texto
        """
        try:
            # Read file content
            content = await file.read()
            
            # Get file extension
            file_extension = file.filename.lower().split('.')[-1]
            file_extension = f'.{file_extension}'
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Formato de archivo no soportado: {file_extension}")
            
            # Process based on file type
            processor = self.supported_formats[file_extension]
            text_content = await processor(content)
            
            logger.info(f"Documento procesado exitosamente: {file.filename}")
            return text_content
            
        except Exception as e:
            logger.error(f"Error procesando documento {file.filename}: {str(e)}")
            raise
    
    async def _process_pdf(self, content: bytes) -> str:
        """
        Procesa archivos PDF
        """
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content += page.extract_text() + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error procesando PDF: {str(e)}")
            raise
    
    async def _process_docx(self, content: bytes) -> str:
        """
        Procesa archivos DOCX
        """
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # También extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error procesando DOCX: {str(e)}")
            raise
    
    async def _process_excel(self, content: bytes) -> str:
        """
        Procesa archivos Excel (XLSX y XLS)
        """
        try:
            excel_file = io.BytesIO(content)
            
            # Intentar con openpyxl primero (para XLSX)
            try:
                workbook = openpyxl.load_workbook(excel_file, data_only=True)
                text_content = ""
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text_content += f"=== HOJA: {sheet_name} ===\n"
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " ".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip():
                            text_content += row_text + "\n"
                    
                    text_content += "\n"
                
                return text_content.strip()
                
            except Exception:
                # Si falla openpyxl, intentar con pandas
                excel_file.seek(0)
                df_list = pd.read_excel(excel_file, sheet_name=None)
                
                text_content = ""
                for sheet_name, df in df_list.items():
                    text_content += f"=== HOJA: {sheet_name} ===\n"
                    text_content += df.to_string(index=False) + "\n\n"
                
                return text_content.strip()
                
        except Exception as e:
            logger.error(f"Error procesando Excel: {str(e)}")
            raise
    
    def extract_tables_from_excel(self, content: bytes) -> Dict[str, pd.DataFrame]:
        """
        Extrae tablas específicas de archivos Excel
        """
        try:
            excel_file = io.BytesIO(content)
            df_list = pd.read_excel(excel_file, sheet_name=None)
            return df_list
            
        except Exception as e:
            logger.error(f"Error extrayendo tablas de Excel: {str(e)}")
            raise 