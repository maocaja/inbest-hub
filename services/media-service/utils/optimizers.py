import os
from PIL import Image
from typing import Tuple, Optional, Dict, Any
import PyPDF2
from docx import Document
from openpyxl import load_workbook
from config import Config

def optimize_image(image_path: str, output_path: str = None) -> Dict[str, Any]:
    """
    Optimiza una imagen reduciendo su tamaño y calidad
    Returns: Dict con información del procesamiento
    """
    try:
        with Image.open(image_path) as img:
            # Convertir a RGB si es necesario
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')
            
            # Obtener dimensiones originales
            original_width, original_height = img.size
            
            # Redimensionar si es muy grande
            if original_width > Config.MAX_IMAGE_WIDTH or original_height > Config.MAX_IMAGE_HEIGHT:
                img.thumbnail((Config.MAX_IMAGE_WIDTH, Config.MAX_IMAGE_HEIGHT), Image.Resampling.LANCZOS)
            
            # Usar el mismo archivo si no se especifica output_path
            if output_path is None:
                output_path = image_path
            
            # Guardar con optimización
            img.save(output_path, quality=Config.IMAGE_QUALITY, optimize=True)
            
            # Obtener tamaño del archivo optimizado
            optimized_size = os.path.getsize(output_path)
            
            return {
                "success": True,
                "original_size": original_width * original_height,
                "optimized_size": img.size[0] * img.size[1],
                "file_size_reduction": os.path.getsize(image_path) - optimized_size,
                "width": img.size[0],
                "height": img.size[1]
            }
            
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

def extract_image_metadata(image_path: str) -> Dict[str, Any]:
    """
    Extrae metadatos de una imagen
    """
    try:
        with Image.open(image_path) as img:
            return {
                "width": img.size[0],
                "height": img.size[1],
                "format": img.format,
                "mode": img.mode,
                "size_bytes": os.path.getsize(image_path)
            }
    except Exception as e:
        return {
            "error": str(e)
        }

def extract_pdf_metadata(pdf_path: str) -> Dict[str, Any]:
    """
    Extrae metadatos de un archivo PDF
    """
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            return {
                "pages": len(pdf_reader.pages),
                "size_bytes": os.path.getsize(pdf_path),
                "title": pdf_reader.metadata.get('/Title', ''),
                "author": pdf_reader.metadata.get('/Author', ''),
                "subject": pdf_reader.metadata.get('/Subject', '')
            }
    except Exception as e:
        return {
            "error": str(e)
        }

def extract_docx_metadata(docx_path: str) -> Dict[str, Any]:
    """
    Extrae metadatos de un archivo DOCX
    """
    try:
        doc = Document(docx_path)
        
        # Contar párrafos y caracteres
        paragraph_count = len(doc.paragraphs)
        character_count = sum(len(paragraph.text) for paragraph in doc.paragraphs)
        
        return {
            "paragraphs": paragraph_count,
            "characters": character_count,
            "size_bytes": os.path.getsize(docx_path)
        }
    except Exception as e:
        return {
            "error": str(e)
        }

def extract_excel_metadata(excel_path: str) -> Dict[str, Any]:
    """
    Extrae metadatos de un archivo Excel
    """
    try:
        workbook = load_workbook(excel_path, read_only=True)
        
        sheet_count = len(workbook.sheetnames)
        total_rows = 0
        total_cells = 0
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            rows = sheet.max_row
            cols = sheet.max_column
            total_rows += rows
            total_cells += rows * cols
        
        return {
            "sheets": sheet_count,
            "total_rows": total_rows,
            "total_cells": total_cells,
            "size_bytes": os.path.getsize(excel_path)
        }
    except Exception as e:
        return {
            "error": str(e)
        }

def extract_file_metadata(file_path: str, file_type: str) -> Dict[str, Any]:
    """
    Extrae metadatos según el tipo de archivo
    """
    if file_type == "image":
        return extract_image_metadata(file_path)
    elif file_type == "document":
        extension = os.path.splitext(file_path)[1].lower()
        
        if extension == '.pdf':
            return extract_pdf_metadata(file_path)
        elif extension == '.docx':
            return extract_docx_metadata(file_path)
        elif extension in ['.xlsx', '.xls']:
            return extract_excel_metadata(file_path)
        else:
            return {
                "size_bytes": os.path.getsize(file_path)
            }
    else:
        return {
            "size_bytes": os.path.getsize(file_path)
        }

def create_thumbnail(image_path: str, thumbnail_path: str, size: Tuple[int, int] = (150, 150)) -> bool:
    """
    Crea una miniatura de una imagen
    """
    try:
        with Image.open(image_path) as img:
            # Convertir a RGB si es necesario
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')
            
            # Crear miniatura manteniendo proporción
            img.thumbnail(size, Image.Resampling.LANCZOS)
            
            # Guardar miniatura
            img.save(thumbnail_path, quality=85, optimize=True)
            
            return True
    except Exception as e:
        print(f"Error creando miniatura: {e}")
        return False 