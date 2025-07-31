"""
Procesador de documentos para extraer información de proyectos inmobiliarios
"""

import os
import logging
from typing import Dict, Any, Optional, List
from pathlib import Path
import PyPDF2
from docx import Document
import openpyxl
from PIL import Image
import io

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Procesador de documentos para extraer información de proyectos inmobiliarios
    """
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image
        }
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Procesar documento y extraer información
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
            
            file_extension = file_path.suffix.lower()
            if file_extension not in self.supported_formats:
                raise ValueError(f"Formato no soportado: {file_extension}")
            
            # Procesar según el tipo de archivo
            processor = self.supported_formats[file_extension]
            extracted_data = processor(file_path)
            
            logger.info(f"Documento procesado exitosamente: {file_path}")
            return {
                "success": True,
                "extracted_data": extracted_data,
                "file_info": {
                    "filename": file_path.name,
                    "size": file_path.stat().st_size,
                    "type": file_extension
                }
            }
            
        except Exception as e:
            logger.error(f"Error procesando documento {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "extracted_data": {}
            }
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """
        Procesar archivo PDF
        """
        extracted_data = {
            "text_content": "",
            "metadata": {},
            "project_info": {}
        }
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extraer texto
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                extracted_data["text_content"] = text_content
                
                # Extraer metadatos
                if pdf_reader.metadata:
                    extracted_data["metadata"] = {
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "subject": pdf_reader.metadata.get('/Subject', ''),
                        "creator": pdf_reader.metadata.get('/Creator', ''),
                        "producer": pdf_reader.metadata.get('/Producer', ''),
                        "creation_date": pdf_reader.metadata.get('/CreationDate', ''),
                        "modification_date": pdf_reader.metadata.get('/ModDate', '')
                    }
                
                # Intentar extraer información del proyecto
                project_info = self._extract_project_info_from_text(text_content)
                extracted_data["project_info"] = project_info
                
        except Exception as e:
            logger.error(f"Error procesando PDF {file_path}: {e}")
            raise
        
        return extracted_data
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """
        Procesar archivo DOCX
        """
        extracted_data = {
            "text_content": "",
            "metadata": {},
            "project_info": {}
        }
        
        try:
            doc = Document(file_path)
            
            # Extraer texto
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"
            
            extracted_data["text_content"] = text_content
            
            # Extraer metadatos básicos
            extracted_data["metadata"] = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            }
            
            # Intentar extraer información del proyecto
            project_info = self._extract_project_info_from_text(text_content)
            extracted_data["project_info"] = project_info
            
        except Exception as e:
            logger.error(f"Error procesando DOCX {file_path}: {e}")
            raise
        
        return extracted_data
    
    def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """
        Procesar archivo Excel
        """
        extracted_data = {
            "sheets": {},
            "metadata": {},
            "project_info": {}
        }
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            # Procesar cada hoja
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []
                
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        sheet_data.append(row)
                
                extracted_data["sheets"][sheet_name] = sheet_data
            
            # Extraer metadatos
            extracted_data["metadata"] = {
                "sheets": workbook.sheetnames,
                "active_sheet": workbook.active.title if workbook.active else None
            }
            
            # Intentar extraer información del proyecto de todas las hojas
            all_text = ""
            for sheet_name, sheet_data in extracted_data["sheets"].items():
                for row in sheet_data:
                    all_text += " ".join(str(cell) for cell in row if cell) + "\n"
            
            project_info = self._extract_project_info_from_text(all_text)
            extracted_data["project_info"] = project_info
            
        except Exception as e:
            logger.error(f"Error procesando Excel {file_path}: {e}")
            raise
        
        return extracted_data
    
    def _process_image(self, file_path: Path) -> Dict[str, Any]:
        """
        Procesar archivo de imagen (placeholder para OCR futuro)
        """
        extracted_data = {
            "image_info": {},
            "project_info": {}
        }
        
        try:
            with Image.open(file_path) as img:
                extracted_data["image_info"] = {
                    "format": img.format,
                    "mode": img.mode,
                    "size": img.size,
                    "width": img.width,
                    "height": img.height
                }
            
            # TODO: Implementar OCR para extraer texto de imágenes
            # Por ahora, solo retornamos información básica de la imagen
            logger.info(f"Imagen procesada: {file_path} (OCR no implementado)")
            
        except Exception as e:
            logger.error(f"Error procesando imagen {file_path}: {e}")
            raise
        
        return extracted_data
    
    def _extract_project_info_from_text(self, text: str) -> Dict[str, Any]:
        """
        Extraer información del proyecto del texto
        """
        project_info = {}
        
        # Convertir a minúsculas para búsqueda
        text_lower = text.lower()
        
        # Buscar patrones comunes de proyectos inmobiliarios
        patterns = {
            "name": [
                r"proyecto[:\s]+([^\n\r]+)",
                r"residencial[:\s]+([^\n\r]+)",
                r"torre[:\s]+([^\n\r]+)",
                r"conjunto[:\s]+([^\n\r]+)",
                r"edificio[:\s]+([^\n\r]+)"
            ],
            "location": [
                r"ubicación[:\s]+([^\n\r]+)",
                r"dirección[:\s]+([^\n\r]+)",
                r"ciudad[:\s]+([^\n\r]+)",
                r"barrio[:\s]+([^\n\r]+)",
                r"zona[:\s]+([^\n\r]+)"
            ],
            "price": [
                r"precio[:\s]+([^\n\r]+)",
                r"valor[:\s]+([^\n\r]+)",
                r"costo[:\s]+([^\n\r]+)",
                r"desde[:\s]+([^\n\r]+)",
                r"hasta[:\s]+([^\n\r]+)"
            ],
            "units": [
                r"apartamentos[:\s]+([^\n\r]+)",
                r"unidades[:\s]+([^\n\r]+)",
                r"pisos[:\s]+([^\n\r]+)",
                r"tipos[:\s]+([^\n\r]+)"
            ],
            "area": [
                r"área[:\s]+([^\n\r]+)",
                r"metros[:\s]+([^\n\r]+)",
                r"m²[:\s]+([^\n\r]+)",
                r"tamaño[:\s]+([^\n\r]+)"
            ],
            "delivery": [
                r"entrega[:\s]+([^\n\r]+)",
                r"fecha[:\s]+([^\n\r]+)",
                r"disponible[:\s]+([^\n\r]+)",
                r"listo[:\s]+([^\n\r]+)"
            ]
        }
        
        import re
        
        for field, pattern_list in patterns.items():
            for pattern in pattern_list:
                matches = re.findall(pattern, text_lower, re.IGNORECASE)
                if matches:
                    project_info[field] = matches[0].strip()
                    break
        
        # Buscar información de contacto
        contact_patterns = {
            "phone": r"(\d{3}[-.\s]?\d{3}[-.\s]?\d{4})",
            "email": r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})",
            "website": r"(https?://[^\s]+)"
        }
        
        for field, pattern in contact_patterns.items():
            matches = re.findall(pattern, text)
            if matches:
                project_info[field] = matches[0]
        
        return project_info
    
    def get_supported_formats(self) -> List[str]:
        """
        Obtener formatos soportados
        """
        return list(self.supported_formats.keys())
    
    def validate_file(self, file_path: str) -> Dict[str, Any]:
        """
        Validar archivo antes de procesar
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return {"valid": False, "error": "Archivo no encontrado"}
            
            if not file_path.is_file():
                return {"valid": False, "error": "No es un archivo válido"}
            
            file_extension = file_path.suffix.lower()
            if file_extension not in self.supported_formats:
                return {
                    "valid": False, 
                    "error": f"Formato no soportado: {file_extension}",
                    "supported_formats": self.get_supported_formats()
                }
            
            file_size = file_path.stat().st_size
            max_size = 10 * 1024 * 1024  # 10MB
            
            if file_size > max_size:
                return {"valid": False, "error": "Archivo demasiado grande"}
            
            return {"valid": True, "file_info": {
                "name": file_path.name,
                "size": file_size,
                "type": file_extension
            }}
            
        except Exception as e:
            return {"valid": False, "error": str(e)} 